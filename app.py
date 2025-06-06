from flask import Flask, render_template, request, Response, jsonify, send_file, send_from_directory, stream_with_context
from flask_socketio import SocketIO, emit
import os
from datetime import datetime
from patent_gen.patent_generator import PatentGenerator
from docx import Document
from io import BytesIO
from patent_gen.agents import get_agent_steps, patent_agent
from patent_gen.events import event_emitter

app = Flask(__name__)
app.config['SECRET_KEY'] = 'your-secret-key'
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 限制上传文件大小为16MB
socketio = SocketIO(app)

# 初始化事件发射器并传入socketio实例
event_emitter.init_app(socketio)

# 全局变量
patent_generator = None
examples = None

@app.route('/')
def index():
    return render_template('index.html')

@socketio.on('connect')
def handle_connect():
    print('Client connected')
    # 发送当前所有步骤
    try:
        steps = get_agent_steps()
        if steps:
            for step in steps:
                socketio.emit('step_update', {'step': step})
    except Exception as e:
        print(f"Error sending steps on connect: {str(e)}")

@socketio.on('disconnect')
def handle_disconnect():
    print('Client disconnected')

@app.route('/upload', methods=['POST'])
def upload_file():
    global patent_generator, examples
    try:
        if 'file' not in request.files:
            return jsonify({'success': False, 'error': '没有文件被上传'})
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({'success': False, 'error': '没有选择文件'})
        
        if not file.filename.endswith('.pdf'):
            return jsonify({'success': False, 'error': '请上传PDF文件'})
        
        # 保存文件
        upload_folder = 'uploads'
        if not os.path.exists(upload_folder):
            os.makedirs(upload_folder)
        
        file_path = os.path.join(upload_folder, file.filename)
        file.save(file_path)
        
        # 初始化PatentGenerator
        patent_generator = PatentGenerator()
        examples = patent_generator.read_pdf_examples(os.path.join(os.path.dirname(__file__), 'examples'))
        print(examples)
        
        if patent_generator.process_pdf_file(file_path):
            return jsonify({
                'success': True,
                'message': '文件上传成功'
            })
        else:
            return jsonify({'success': False, 'error': 'PDF处理失败'})

    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})

@app.route('/generate', methods=['POST'])
def generate_patent():
    global patent_generator, examples
    try:
        if not patent_generator.content:
            return jsonify({'success': False, 'error': '请先上传PDF文件'})

        data = request.get_json()
        step = data.get('step')

        if step == 'abstract':
            abstract_parts = ""
            for abstract_part in patent_generator.generate_patent_abstract(examples):
                print(abstract_part)
                abstract_parts += abstract_part
            return jsonify({'success': True, 'content': abstract_parts})
        

        elif step == 'claims':
            claims_parts = ""
            for claims_part in patent_generator.generate_patent_claims(examples):
                print(claims_part)
                claims_parts += claims_part
            return jsonify({'success': True, 'content': claims_parts})


        elif step == 'description':
            description_parts = ""
            for description_part in patent_generator.generate_patent_description(examples):
                print(description_part)
                description_parts += description_part
            return jsonify({'success': True, 'content': description_parts})
        else:
            return jsonify({'success': True, 'content': "else"})
            

    except Exception as e:
        socketio.emit('step_update', {'message': f'生成失败: {str(e)}'}, namespace='/test')
        return jsonify({'success': False, 'error': str(e)})

@app.route('/export', methods=['POST'])
def export_document():
    try:
        data = request.get_json()
        if not data:
            return jsonify({'success': False, 'error': '没有接收到数据'})
        
        # 创建Word文档
        doc = Document()
        
        # 添加标题
        doc.add_heading('专利申请文档', 0)
        
        # 添加摘要
        doc.add_heading('摘要', level=1)
        doc.add_paragraph(data.get('abstract', ''))
        
        # 添加权利要求
        doc.add_heading('权利要求', level=1)
        doc.add_paragraph(data.get('claims', ''))
        
        # 添加说明书
        doc.add_heading('说明书', level=1)
        doc.add_paragraph(data.get('description', ''))
        
        # 保存到内存中
        doc_io = BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)
        
        # 生成文件名
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f'patent_{timestamp}.docx'
        
        return send_file(
            doc_io,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=filename
        )
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})

@app.route('/get_steps')
def get_steps():
    try:
        steps_data = get_agent_steps()
        return jsonify({
            'success': True,
            'steps': steps_data
        })
    except Exception as e:
        return jsonify({
            'success': False,
            'error': str(e)
        })

if __name__ == '__main__':
    socketio.run(app, debug=True)
