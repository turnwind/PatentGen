import os
import PyPDF2
import docx
import tkinter as tk
from tkinter import ttk, filedialog, scrolledtext, messagebox
from openai import OpenAI
from config import OPENAI_API_KEY, OPENAI_MODEL, INPUT_DIR, OUTPUT_DIR, OPENAI_BASE_URL
from patent_generator import PatentGenerator
from events import event_emitter

class PatentGeneratorGUI:
    def __init__(self, root):
        self.root = root
        self.root.title("专利生成器")
        self.root.geometry("800x600")
        
        # 初始化专利生成器
        self.patent_generator = None
        
        # 创建主框架
        self.main_frame = ttk.Frame(root, padding="10")
        self.main_frame.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        
        # 创建文件选择部分
        self.create_file_selection()
        
        # 创建进度显示部分
        self.create_progress_section()
        
        # 创建日志显示部分
        self.create_log_section()
        
        # 确保目录存在
        for dir_path in [INPUT_DIR, OUTPUT_DIR, 'uploads']:
            os.makedirs(dir_path, exist_ok=True)
        
        # 设置事件处理
        event_emitter.on_step_update = self.handle_step_update

    def create_file_selection(self):
        # 文件选择框架
        file_frame = ttk.LabelFrame(self.main_frame, text="选择论文", padding="5")
        file_frame.grid(row=0, column=0, columnspan=2, sticky=(tk.W, tk.E), pady=5)
        
        # 文件路径显示
        self.file_path_var = tk.StringVar()
        self.file_path_entry = ttk.Entry(file_frame, textvariable=self.file_path_var, width=60)
        self.file_path_entry.grid(row=0, column=0, padx=5)
        
        # 浏览按钮
        browse_btn = ttk.Button(file_frame, text="浏览", command=self.browse_file)
        browse_btn.grid(row=0, column=1, padx=5)
        
        # 生成按钮
        self.generate_btn = ttk.Button(file_frame, text="生成专利", command=self.generate_patent)
        self.generate_btn.grid(row=0, column=2, padx=5)
        self.generate_btn['state'] = 'disabled'

    def create_progress_section(self):
        # 进度框架
        progress_frame = ttk.LabelFrame(self.main_frame, text="生成进度", padding="5")
        progress_frame.grid(row=1, column=0, columnspan=2, sticky=(tk.W, tk.E), pady=5)
        
        # 进度条
        self.progress_var = tk.DoubleVar()
        self.progress_bar = ttk.Progressbar(
            progress_frame, 
            variable=self.progress_var,
            maximum=100,
            length=780
        )
        self.progress_bar.grid(row=0, column=0, padx=5, pady=5)

    def create_log_section(self):
        # 日志框架
        log_frame = ttk.LabelFrame(self.main_frame, text="处理日志", padding="5")
        log_frame.grid(row=2, column=0, columnspan=2, sticky=(tk.W, tk.E, tk.N, tk.S), pady=5)
        
        # 日志文本框
        self.log_text = scrolledtext.ScrolledText(log_frame, width=90, height=20)
        self.log_text.grid(row=0, column=0, padx=5, pady=5)

    def browse_file(self):
        filename = filedialog.askopenfilename(
            title="选择论文PDF文件",
            filetypes=[("PDF files", "*.pdf")]
        )
        if filename:
            self.file_path_var.set(filename)
            self.generate_btn['state'] = 'normal'

    def log_message(self, message):
        self.log_text.insert(tk.END, message + "\n")
        self.log_text.see(tk.END)
        self.root.update()

    def handle_step_update(self, step_data):
        """处理步骤更新事件"""
        if not step_data:
            return
            
        # 更新日志
        if step_data.get('message'):
            self.log_message(step_data['message'])
        
        # 更新进度
        if step_data.get('status') == 'completed':
            current = self.progress_var.get()
            self.progress_var.set(min(current + 20, 100))
        
        # 处理错误
        if step_data.get('status') == 'error':
            messagebox.showerror("错误", step_data.get('message', '生成过程出错'))
            self.generate_btn['state'] = 'normal'
            self.progress_var.set(0)

    def generate_patent(self):
        input_file = self.file_path_var.get()
        if not input_file:
            messagebox.showerror("错误", "请选择论文PDF文件")
            return
        
        try:
            # 禁用生成按钮
            self.generate_btn['state'] = 'disabled'
            self.progress_var.set(0)
            self.log_message(f"开始处理文件: {input_file}")
            
            # 初始化专利生成器
            self.patent_generator = PatentGenerator()
            
            # 处理PDF文件
            if not self.patent_generator.process_pdf_file(input_file):
                raise Exception("PDF处理失败")
            
            # 生成专利文档
            result = self.patent_generator.generate_patent()
            if not result:
                raise Exception("生成专利文档失败")
            
            # 导出Word文档
            output_path = os.path.join(
                OUTPUT_DIR,
                f"专利申请书_{os.path.splitext(os.path.basename(input_file))[0]}.docx"
            )
            
            doc = docx.Document()
            doc.add_heading('专利申请文档', 0)
            
            doc.add_heading('摘要', level=1)
            doc.add_paragraph(result['abstract'])
            
            doc.add_heading('权利要求', level=1)
            doc.add_paragraph(result['claims'])
            
            doc.add_heading('说明书', level=1)
            doc.add_paragraph(result['description'])
            
            doc.save(output_path)
            
            self.log_message(f"专利文档已生成: {output_path}")
            messagebox.showinfo("成功", "专利文档生成完成！")
            
        except Exception as e:
            self.log_message(f"错误: {str(e)}")
            messagebox.showerror("错误", str(e))
        finally:
            self.generate_btn['state'] = 'normal'

def main():
    root = tk.Tk()
    app = PatentGeneratorGUI(root)
    root.mainloop()

if __name__ == "__main__":
    main()
