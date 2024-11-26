import os
import PyPDF2
import docx
from openai import OpenAI
from config import OPENAI_API_KEY, OPENAI_MODEL, INPUT_DIR, OUTPUT_DIR, PATENT_PROMPTS, OPENAI_BASE_URL
from agents import AnalysisAgent, AbstractAgent, ClaimsAgent, DescriptionAgent, DrawingDescriptionAgent

class PatentGenerator:
    def __init__(self):
        self.client = OpenAI(api_key=OPENAI_API_KEY, base_url=OPENAI_BASE_URL)
        self.analysis_agent = AnalysisAgent()
        self.abstract_agent = AbstractAgent()
        self.claims_agent = ClaimsAgent()
        self.description_agent = DescriptionAgent()
        self.drawing_agent = DrawingDescriptionAgent()
        
    def extract_text_from_pdf(self, pdf_path):
        """从PDF文件中提取文本"""
        try:
            with open(pdf_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                text = ""
                for page in reader.pages:
                    text += page.extract_text()
                return text
        except Exception as e:
            print(f"Error extracting text from PDF: {e}")
            return None

    def generate_patent_content(self, text, prompt_key):
        """使用OpenAI API生成专利内容"""
        try:
            # 首先进行技术分析
            analysis_report = self.analysis_agent.process(text)
            if not analysis_report:
                return None

            # 根据不同部分选择不同的Agent
            if prompt_key == "abstract":
                return self.abstract_agent.process(analysis_report)
            elif prompt_key == "claims":
                return self.claims_agent.process(analysis_report)
            elif prompt_key == "description":
                description = self.description_agent.process(analysis_report)
                drawing_description = self.drawing_agent.process(analysis_report)
                return f"{description}\n\n{drawing_description}"
            
        except Exception as e:
            print(f"Error generating patent content: {e}")
            return None

    def create_patent_document(self, abstract, claims, description, output_path):
        """创建专利文档"""
        try:
            doc = docx.Document()
            
            # 添加标题
            doc.add_heading('发明专利申请', 0)
            
            # 添加摘要
            doc.add_heading('摘要', 1)
            doc.add_paragraph(abstract)
            
            # 添加权利要求
            doc.add_heading('权利要求', 1)
            doc.add_paragraph(claims)
            
            # 添加说明书
            doc.add_heading('说明书', 1)
            doc.add_paragraph(description)
            
            # 保存文档
            doc.save(output_path)
            return True
        except Exception as e:
            print(f"Error creating patent document: {e}")
            return False

    def process_paper(self, input_file):
        """处理论文并生成专利文档"""
        # 检查文件是否存在
        if not os.path.exists(input_file):
            print(f"Input file not found: {input_file}")
            return False

        # 提取文本
        paper_text = self.extract_text_from_pdf(input_file)
        if not paper_text:
            return False

        # 生成专利内容
        print("正在生成专利摘要...")
        abstract = self.generate_patent_content(paper_text, "abstract")
        
        print("正在生成权利要求...")
        claims = self.generate_patent_content(paper_text, "claims")
        
        print("正在生成说明书...")
        description = self.generate_patent_content(paper_text, "description")

        if not all([abstract, claims, description]):
            print("Error generating patent content")
            return False

        # 创建输出文件名
        output_file = os.path.join(
            OUTPUT_DIR,
            f"专利申请书_{os.path.splitext(os.path.basename(input_file))[0]}.docx"
        )

        # 创建专利文档
        if self.create_patent_document(abstract, claims, description, output_file):
            print(f"专利文档已生成: {output_file}")
            return True
        return False

def main():
    # 检查必要的目录是否存在
    if not os.path.exists(INPUT_DIR):
        print(f"Input directory not found: {INPUT_DIR}")
        return

    if not os.path.exists(OUTPUT_DIR):
        os.makedirs(OUTPUT_DIR)

    # 检查API密钥
    if not OPENAI_API_KEY:
        print("请在.env文件中设置OPENAI_API_KEY")
        return

    generator = PatentGenerator()

    # 处理input目录中的所有PDF文件
    pdf_files = [f for f in os.listdir(INPUT_DIR) if f.lower().endswith('.pdf')]
    
    if not pdf_files:
        print(f"在{INPUT_DIR}目录中没有找到PDF文件")
        return

    for pdf_file in pdf_files:
        input_path = os.path.join(INPUT_DIR, pdf_file)
        print(f"\n处理文件: {pdf_file}")
        generator.process_paper(input_path)

if __name__ == "__main__":
    main()
